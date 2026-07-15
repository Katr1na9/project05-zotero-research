from __future__ import annotations

import hashlib
import json
import sys
from collections import Counter
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn


def points(value):
    return None if value is None else round(value.pt, 2)


def inches(value):
    return None if value is None else round(value.inches, 3)


def paragraph_record(index, paragraph):
    fmt = paragraph.paragraph_format
    run_fonts = []
    for run in paragraph.runs:
        if not run.text.strip():
            continue
        rpr = run._element.rPr
        east_asia = None
        if rpr is not None and rpr.rFonts is not None:
            east_asia = rpr.rFonts.get(qn("w:eastAsia"))
        run_fonts.append(
            {
                "text_length": len(run.text),
                "font": run.font.name,
                "east_asia": east_asia,
                "size_pt": points(run.font.size),
                "bold": run.bold,
            }
        )
    return {
        "index": index,
        "text": paragraph.text,
        "style": paragraph.style.name,
        "alignment": str(paragraph.alignment),
        "left_indent_in": inches(fmt.left_indent),
        "right_indent_in": inches(fmt.right_indent),
        "first_line_indent_in": inches(fmt.first_line_indent),
        "space_before_pt": points(fmt.space_before),
        "space_after_pt": points(fmt.space_after),
        "line_spacing": str(fmt.line_spacing),
        "keep_with_next": fmt.keep_with_next,
        "page_break_before": fmt.page_break_before,
        "runs": run_fonts,
    }


def analyze(path: Path):
    doc = Document(path)
    records = [paragraph_record(i, p) for i, p in enumerate(doc.paragraphs)]
    with ZipFile(path) as zf:
        names = zf.namelist()
        document_xml = zf.read("word/document.xml")
        comments = zf.read("word/comments.xml") if "word/comments.xml" in names else b""
        package = [
            {
                "path": name,
                "size": zf.getinfo(name).file_size,
                "sha256": hashlib.sha256(zf.read(name)).hexdigest(),
            }
            for name in names
        ]
    sections = []
    for i, section in enumerate(doc.sections, start=1):
        sections.append(
            {
                "number": i,
                "width_in": inches(section.page_width),
                "height_in": inches(section.page_height),
                "top_in": inches(section.top_margin),
                "bottom_in": inches(section.bottom_margin),
                "left_in": inches(section.left_margin),
                "right_in": inches(section.right_margin),
                "header_distance_in": inches(section.header_distance),
                "footer_distance_in": inches(section.footer_distance),
                "header": [p.text for p in section.header.paragraphs],
                "footer": [p.text for p in section.footer.paragraphs],
            }
        )
    used_styles = Counter(record["style"] for record in records)
    return {
        "path": str(path),
        "sha256": hashlib.sha256(path.read_bytes()).hexdigest(),
        "paragraph_count": len(records),
        "table_count": len(doc.tables),
        "inline_shape_count": len(doc.inline_shapes),
        "section_count": len(doc.sections),
        "tracked_insertions": document_xml.count(b"<w:ins"),
        "tracked_deletions": document_xml.count(b"<w:del"),
        "comment_count": comments.count(b"<w:comment "),
        "used_styles": used_styles,
        "sections": sections,
        "paragraphs": records,
        "package": package,
    }


def main():
    output = Path(sys.argv[1])
    result = [analyze(Path(item)) for item in sys.argv[2:]]
    output.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    for item in result:
        print(Path(item["path"]).name)
        print(
            f"paragraphs={item['paragraph_count']} sections={item['section_count']} "
            f"tables={item['table_count']} shapes={item['inline_shape_count']} "
            f"ins={item['tracked_insertions']} del={item['tracked_deletions']} comments={item['comment_count']}"
        )
        for record in item["paragraphs"]:
            if record["text"].strip():
                print(f"{record['index']:03d}\t{record['style']}\t{record['alignment']}\t{record['text'][:120]}")


if __name__ == "__main__":
    main()
