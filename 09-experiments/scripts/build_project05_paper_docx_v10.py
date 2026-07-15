#!/usr/bin/env python3
"""Build the Project05 Chinese paper working manuscript as a deterministic DOCX."""

from __future__ import annotations

import argparse
import hashlib
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
DEFAULT_MD = ROOT / "08-writing" / "paper-main-draft-v1.0-parameter-governance-20260715.md"
DEFAULT_BIB = ROOT / "08-writing" / "paper-main-references-v0.3.bib"
DEFAULT_FIGURE = ROOT / "08-writing" / "figures" / "main-v0.4" / "fig1_method_and_information_boundary.png"
DEFAULT_OUT = ROOT / "08-writing" / "paper-package-v1.0-parameter-governance" / "Project05_不完整证据下的可审计APT调查控制-完整工作稿.docx"


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as fh:
        for block in iter(lambda: fh.read(1024 * 1024), b""):
            h.update(block)
    return h.hexdigest().upper()


def clean_latex(value: str) -> str:
    value = value.replace("\\&", "&").replace("~", " ")
    value = value.replace("---", "—").replace("--", "–")
    value = re.sub(r"\\[a-zA-Z]+\s*", "", value)
    value = value.replace("{", "").replace("}", "")
    return re.sub(r"\s+", " ", value).strip()


def latex_to_plain(value: str) -> str:
    """Convert the small LaTeX subset used in the manuscript to readable Word text."""
    replacements = (
        (r"\operatorname{Cover}", "Cover"),
        (r"\operatorname{VoI}", "VoI"),
        (r"\mathcal{C}", "C"),
        (r"\mathbb{E}", "E"),
        (r"\Longleftrightarrow", "⇔"),
        (r"\ge", "≥"),
        (r"\le", "≤"),
        (r"\ne", "≠"),
        (r"\in", "∈"),
        (r"\mid", "|"),
        (r"\max", "max"),
        (r"\alpha", "α"),
        (r"\star", "*"),
        (r"\left", ""),
        (r"\right", ""),
        (r"\,", " "),
        (r"\{", "<<LBRACE>>"),
        (r"\}", "<<RBRACE>>"),
    )
    for old, new in replacements:
        value = value.replace(old, new)
    value = re.sub(r"\^\{([^{}]+)\}", r"^\1", value)
    value = re.sub(r"_\{([^{}]+)\}", r"_(\1)", value)
    value = value.replace("^{", "^").replace("_{", "_")
    value = value.replace("}", "").replace("{", "")
    value = value.replace("<<LBRACE>>", "{").replace("<<RBRACE>>", "}")
    value = re.sub(r"\\([A-Za-z]+)", r"\1", value)
    value = re.sub(r"\s+", " ", value)
    return value.strip()


def plain_heading_text(value: str) -> str:
    return re.sub(r"\$([^$]+)\$", lambda m: latex_to_plain(m.group(1)), value)


def parse_bib(path: Path) -> dict[str, dict[str, str]]:
    text = path.read_text(encoding="utf-8")
    entries: dict[str, dict[str, str]] = {}
    start_re = re.compile(r"@(\w+)\s*\{\s*([^,]+),", re.M)
    for match in start_re.finditer(text):
        depth = 1
        i = match.end()
        quoted = False
        escaped = False
        while i < len(text) and depth:
            ch = text[i]
            if escaped:
                escaped = False
            elif ch == "\\":
                escaped = True
            elif ch == '"':
                quoted = not quoted
            elif not quoted and ch == "{":
                depth += 1
            elif not quoted and ch == "}":
                depth -= 1
            i += 1
        body = text[match.end() : i - 1]
        fields: dict[str, str] = {"ENTRYTYPE": match.group(1).lower()}
        j = 0
        while j < len(body):
            fm = re.search(r"([A-Za-z][A-Za-z0-9_-]*)\s*=\s*", body[j:])
            if not fm:
                break
            name = fm.group(1).lower()
            pos = j + fm.end()
            if pos >= len(body):
                break
            opener = body[pos]
            if opener == "{":
                k, d = pos + 1, 1
                quoted2 = False
                esc2 = False
                while k < len(body) and d:
                    ch = body[k]
                    if esc2:
                        esc2 = False
                    elif ch == "\\":
                        esc2 = True
                    elif ch == '"':
                        quoted2 = not quoted2
                    elif not quoted2 and ch == "{":
                        d += 1
                    elif not quoted2 and ch == "}":
                        d -= 1
                    k += 1
                value = body[pos + 1 : k - 1]
            elif opener == '"':
                k = pos + 1
                esc2 = False
                while k < len(body):
                    ch = body[k]
                    if esc2:
                        esc2 = False
                    elif ch == "\\":
                        esc2 = True
                    elif ch == '"':
                        k += 1
                        break
                    k += 1
                value = body[pos + 1 : k - 1]
            else:
                k = pos
                while k < len(body) and body[k] not in ",\n":
                    k += 1
                value = body[pos:k]
            fields[name] = clean_latex(value)
            j = k
        entries[match.group(2).strip()] = fields
    return entries


def replace_citations(text: str) -> tuple[str, list[str]]:
    order: list[str] = []

    def repl(match: re.Match[str]) -> str:
        keys = re.findall(r"@([A-Za-z0-9_:\-]+)", match.group(1))
        nums = []
        for key in keys:
            if key not in order:
                order.append(key)
            nums.append(order.index(key) + 1)
        return "[" + ", ".join(str(n) for n in nums) + "]"

    return re.sub(r"\[([^\]]*@[^\]]+)\]", repl, text), order


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def make_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def new_numbering_instance(doc: Document, style_name="List Number") -> int:
    style = doc.styles[style_name]
    base_num_id = int(style._element.pPr.numPr.numId.val)
    numbering = doc.part.numbering_part.element
    base_num = next(
        node
        for node in numbering.findall(qn("w:num"))
        if int(node.get(qn("w:numId"))) == base_num_id
    )
    abstract_id = base_num.find(qn("w:abstractNumId")).get(qn("w:val"))
    new_id = max(int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_id)
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return new_id


def set_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(qn("w:val"), "0")
    num = num_pr.find(qn("w:numId"))
    if num is None:
        num = OxmlElement("w:numId")
        num_pr.append(num)
    num.set(qn("w:val"), str(num_id))


def set_run_font(run, size: float, bold=False, italic=False, name="Times New Roman", east_asia="宋体") -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)


INLINE_RE = re.compile(r"(\*\*.+?\*\*|`.+?`|\$[^$]+?\$)")


def add_inline(paragraph, text: str, size=10.5) -> None:
    pos = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > pos:
            set_run_font(paragraph.add_run(text[pos : match.start()]), size)
        token = match.group(0)
        if token.startswith("**"):
            set_run_font(paragraph.add_run(token[2:-2]), size, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size - 0.5, name="Consolas", east_asia="等线")
            run.font.color.rgb = RGBColor(80, 80, 80)
        else:
            set_run_font(paragraph.add_run(latex_to_plain(token[1:-1])), size, italic=True, name="Cambria Math")
        pos = match.end()
    if pos < len(text):
        set_run_font(paragraph.add_run(text[pos:]), size)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Cm(0.74)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(0)

    if "Paper Title" not in styles:
        title = styles.add_style("Paper Title", WD_STYLE_TYPE.PARAGRAPH)
    else:
        title = styles["Paper Title"]
    title.font.name = "Times New Roman"
    title.font.size = Pt(18)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0, 0, 0)
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)

    for name, size, before, after in (("Heading 1", 14, 12, 6), ("Heading 2", 12, 9, 4), ("Heading 3", 11, 6, 3)):
        style = styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Paper Caption" not in styles:
        cap = styles.add_style("Paper Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        cap = styles["Paper Caption"]
    cap.font.name = "Times New Roman"
    cap.font.size = Pt(9)
    cap._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(3)
    cap.paragraph_format.space_after = Pt(6)
    cap.paragraph_format.keep_with_next = False

    if "Paper Reference" not in styles:
        ref = styles.add_style("Paper Reference", WD_STYLE_TYPE.PARAGRAPH)
    else:
        ref = styles["Paper Reference"]
    ref.font.name = "Times New Roman"
    ref.font.size = Pt(9)
    ref._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    ref.paragraph_format.left_indent = Cm(0.74)
    ref.paragraph_format.first_line_indent = Cm(-0.74)
    ref.paragraph_format.line_spacing = 1.0
    ref.paragraph_format.space_after = Pt(3)


def add_metadata_block(doc: Document, label: str, value: str, bold_value=False) -> None:
    p = doc.add_paragraph()
    centered = label.startswith("英文题目") or label.startswith("作者与单位")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0)
    size = 10.5 if centered else 9.5
    set_run_font(p.add_run(label), size, bold=True, east_asia="黑体")
    set_run_font(p.add_run(value), size, bold=bold_value)


def add_table(doc: Document, rows: list[list[str]], total_width=9020) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    rows = [row + [""] * (cols - len(row)) for row in rows]
    weights = []
    for col in range(cols):
        longest = max(len(re.sub(r"[`*$]", "", row[col])) for row in rows)
        weights.append(max(5, min(28, longest)))
    raw = [max(650, round(total_width * w / sum(weights))) for w in weights]
    raw[-1] += total_width - sum(raw)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(0)
            add_inline(p, value, 8.5)
            if i == 0:
                for run in p.runs:
                    run.bold = True
                set_cell_shading(cell, "E7E6E6")
        if i == 0:
            tr_pr = table.rows[0]._tr.get_or_add_trPr()
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            tr_pr.append(repeat)
    set_table_widths(table, raw)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(1)


def format_reference(number: int, key: str, entry: dict[str, str]) -> str:
    authors = entry.get("author", "[author unavailable]")
    authors = "; ".join(part.strip() for part in authors.split(" and "))
    year = entry.get("year", "n.d.")
    title = entry.get("title", "[title unavailable]")
    venue = entry.get("journal") or entry.get("booktitle") or entry.get("institution") or entry.get("publisher", "")
    details = []
    if venue:
        details.append(venue)
    if entry.get("volume"):
        volume = entry["volume"]
        if entry.get("number"):
            volume += f"({entry['number']})"
        details.append(volume)
    if entry.get("pages"):
        details.append(f"pp. {entry['pages']}")
    locator = f"doi:{entry['doi']}" if entry.get("doi") else entry.get("url", "")
    tail = ", ".join(details)
    if tail:
        tail = f" {tail}."
    if locator:
        tail += f" {locator}"
    return f"[{number}] {authors}. ({year}). {title}.{tail}"


def build(md_path: Path, bib_path: Path, figure_path: Path, out_path: Path) -> dict:
    source = md_path.read_text(encoding="utf-8")
    source = re.sub(r"^---\s*\n.*?\n---\s*\n", "", source, count=1, flags=re.S)
    source, citation_order = replace_citations(source)
    bib = parse_bib(bib_path)
    missing = [key for key in citation_order if key not in bib]
    if missing:
        raise ValueError(f"Missing bibliography keys: {missing}")

    lines = source.splitlines()
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.3)
    section.bottom_margin = Cm(2.3)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.2)
    configure_styles(doc)
    make_page_number(section.footer.paragraphs[0])
    doc.core_properties.title = "不完整证据下的可审计 APT 调查控制：信息边界、参数治理与证据获取"
    doc.core_properties.subject = "Project05 参数治理整合论文完整工作稿"
    doc.core_properties.author = "待作者确认"
    doc.core_properties.keywords = "APT调查, 不完整证据, 主动取证, 信息边界, 参数治理"

    table_buffer: list[list[str]] = []
    in_equation = False
    equation_lines: list[str] = []
    title_done = False
    skip_after_references = False
    figure_inserted = False
    active_num_id: int | None = None
    i = 0
    while i < len(lines):
        raw = lines[i].rstrip()
        stripped = raw.strip()

        if stripped == "## 投稿前必须完成":
            break
        if stripped == "## 参考文献":
            if table_buffer:
                add_table(doc, table_buffer)
                table_buffer = []
            doc.add_paragraph("参考文献", style="Heading 1")
            for num, key in enumerate(citation_order, 1):
                p = doc.add_paragraph(style="Paper Reference")
                add_inline(p, format_reference(num, key, bib[key]), 9)
            skip_after_references = True
            active_num_id = None
            i += 1
            continue
        if skip_after_references:
            i += 1
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            cells = [cell.strip() for cell in stripped.strip("|").split("|")]
            if all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
                i += 1
                continue
            table_buffer.append(cells)
            i += 1
            continue
        if table_buffer:
            add_table(doc, table_buffer)
            table_buffer = []

        if stripped == r"\[":
            in_equation = True
            equation_lines = []
            i += 1
            continue
        if stripped == r"\]":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            set_run_font(p.add_run(latex_to_plain(" ".join(equation_lines))), 10.5, italic=True, name="Cambria Math")
            in_equation = False
            i += 1
            continue
        if in_equation:
            equation_lines.append(stripped)
            i += 1
            continue
        if not stripped:
            active_num_id = None
            i += 1
            continue

        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2)
            if level == 1 and not title_done:
                p = doc.add_paragraph(style="Paper Title")
                add_inline(p, text, 18)
                title_done = True
            else:
                doc.add_paragraph(plain_heading_text(text), style=f"Heading {min(level - 1, 3)}")
            active_num_id = None
            i += 1
            continue

        meta = re.match(r"^\*\*(英文题目|作者与单位|稿件状态|关键词)\*\*[:：]\s*(.*)$", stripped)
        if meta:
            add_metadata_block(doc, meta.group(1) + "：", meta.group(2), bold_value=False)
            active_num_id = None
            i += 1
            continue

        numbered = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        bullet = re.match(r"^-\s+(.+)$", stripped)
        if numbered:
            if active_num_id is None:
                active_num_id = new_numbering_instance(doc)
            p = doc.add_paragraph(style="List Number")
            set_numbering(p, active_num_id)
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            add_inline(p, numbered.group(2))
        elif bullet:
            active_num_id = None
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            add_inline(p, bullet.group(1))
        else:
            active_num_id = None
            p = doc.add_paragraph()
            add_inline(p, stripped)

        if not figure_inserted and "关键在生成过程是否读取隐藏结果" in stripped and figure_path.exists():
            fig_p = doc.add_paragraph()
            fig_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fig_p.paragraph_format.first_line_indent = Cm(0)
            fig_p.add_run().add_picture(str(figure_path), width=Inches(6.15))
            cap = doc.add_paragraph(style="Paper Caption")
            add_inline(cap, "图1 调查控制闭环与信息边界。规划器只访问公开动作目标、成本、当前缺口、预算及历史反馈；执行器和 Oracle 隐藏实际恢复集合与实现通道状态。", 9)
            figure_inserted = True
        i += 1

    if table_buffer:
        add_table(doc, table_buffer)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    audit = {
        "source_markdown": str(md_path.relative_to(ROOT)).replace("\\", "/"),
        "source_bibliography": str(bib_path.relative_to(ROOT)).replace("\\", "/"),
        "output_docx": str(out_path.relative_to(ROOT)).replace("\\", "/"),
        "source_markdown_sha256": sha256(md_path),
        "source_bibliography_sha256": sha256(bib_path),
        "output_docx_sha256": sha256(out_path),
        "citation_keys": citation_order,
        "citation_key_count": len(citation_order),
        "missing_citation_keys": missing,
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "inline_shapes": len(doc.inline_shapes),
        "sections": len(doc.sections),
        "excluded_internal_section": "投稿前必须完成",
        "figure_1_source": str(figure_path.relative_to(ROOT)).replace("\\", "/") if figure_path.exists() else None,
    }
    audit_path = out_path.parent / "build-audit-paper-v1.0.json"
    audit_path.write_text(json.dumps(audit, ensure_ascii=False, indent=2), encoding="utf-8")
    return audit


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--markdown", type=Path, default=DEFAULT_MD)
    parser.add_argument("--bibliography", type=Path, default=DEFAULT_BIB)
    parser.add_argument("--figure", type=Path, default=DEFAULT_FIGURE)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUT)
    args = parser.parse_args()
    audit = build(args.markdown.resolve(), args.bibliography.resolve(), args.figure.resolve(), args.output.resolve())
    print(json.dumps(audit, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
