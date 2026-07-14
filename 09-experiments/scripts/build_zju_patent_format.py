from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from PIL import Image


TITLE_SIZE = Pt(14)
BODY_SIZE = Pt(14)
HEADER_SIZE = Pt(9)
FIRST_LINE = Cm(0.99)


def set_run_font(run, size=BODY_SIZE, bold=None):
    run.font.name = "Times New Roman"
    run.font.size = size
    if bold is not None:
        run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "宋体")


def set_paragraph_format(paragraph, *, center=False, bold=False, indent=True):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    fmt = paragraph.paragraph_format
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.first_line_indent = FIRST_LINE if indent else Cm(0)
    fmt.keep_with_next = True if bold else None
    for run in paragraph.runs:
        set_run_font(run, bold=bold)
    return paragraph


def add_text_paragraph(doc, text, *, center=False, bold=False, indent=True):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run, bold=bold)
    return set_paragraph_format(paragraph, center=center, bold=bold, indent=indent)


def set_text_paragraph(paragraph, text, *, center=False, bold=False, indent=True):
    paragraph.clear()
    run = paragraph.add_run(text)
    set_run_font(run, bold=bold)
    return set_paragraph_format(paragraph, center=center, bold=bold, indent=indent)


def add_header_border(paragraph):
    ppr = paragraph._p.get_or_add_pPr()
    existing = ppr.find(qn("w:pBdr"))
    if existing is not None:
        ppr.remove(existing)
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:color"), "auto")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    borders.append(bottom)
    ppr.append(borders)


def configure_section(section, header_text):
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.175)
    section.right_margin = Cm(3.175)
    section.header_distance = Cm(1.50)
    section.footer_distance = Cm(1.75)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(header_text)
    set_run_font(run, size=HEADER_SIZE, bold=False)
    add_header_border(paragraph)
    footer = section.footer
    footer.paragraphs[0].clear()


def add_new_section(doc, header_text):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section, header_text)
    return section


def split_claim_one(text):
    marker = "包括： "
    if marker not in text:
        return ["1." + text]
    prefix, rest = text.split(marker, 1)
    clauses = re.split(r"；\s+", rest)
    output = ["1." + prefix + "包括："]
    for index, clause in enumerate(clauses, start=1):
        clause = clause.strip()
        if clause.startswith("以及 "):
            clause = clause[3:]
        punctuation = "；" if index < len(clauses) else ""
        output.append(f"S{index}、{clause}{punctuation}")
    return output


def method_steps(claim_one):
    steps = split_claim_one(claim_one)[1:]
    return [step[:-1] if step.endswith("；") else step for step in steps]


def add_claims(doc, claims):
    for claim in claims:
        number = claim["number"]
        text = claim["text"].strip()
        if number == 1:
            for paragraph_text in split_claim_one(text):
                add_text_paragraph(doc, paragraph_text)
        else:
            add_text_paragraph(doc, f"{number}.{text}")


def add_equation(doc, expression):
    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph, center=True, indent=False)
    math_para = OxmlElement("m:oMathPara")
    math_object = OxmlElement("m:oMath")
    math_run = OxmlElement("m:r")
    math_text = OxmlElement("m:t")
    math_text.text = expression
    math_run.append(math_text)
    math_object.append(math_run)
    math_para.append(math_object)
    paragraph._p.append(math_para)


def add_beneficial_effects(doc, blocks):
    if not blocks:
        return
    add_text_paragraph(doc, blocks[0])
    if len(blocks) > 1:
        items = re.split(r"\s+(?=\d+\.\s)", blocks[1].strip())
        for item in items:
            add_text_paragraph(doc, item)
    for block in blocks[2:]:
        add_text_paragraph(doc, block)


def embodiment_lead(heading):
    mapping = {
        "5.1": "在S1输入与数据对象定义中，具体内容如下：",
        "5.2": "在S2证据缺口状态构建中，具体内容如下：",
        "5.3": "在S3信息边界校验与候选动作管理中，具体内容如下：",
        "5.4": "在S4动作评价与选择中，具体内容如下：",
        "5.5": "在S5通道执行与反馈更新中，具体内容如下：",
        "5.6": "在S6停止与降级输出中，具体内容如下：",
        "5.7": "进一步地，在存在动作先决条件或解锁关系时，采用非短视规划实施方式：",
        "5.8": "在真实安全事件实施例中，具体实施结果如下：",
        "5.9": "在多源证据来源核验实施例中，具体处理如下：",
        "5.10": "在可选语义编译实施例中，具体处理如下：",
    }
    prefix = heading.split(maxsplit=1)[0]
    return mapping.get(prefix, heading)


def add_specification(doc, data):
    spec = data["specification"]
    add_text_paragraph(doc, "技术领域", bold=True, indent=False)
    for paragraph in spec["technical_field"]:
        add_text_paragraph(doc, paragraph)

    add_text_paragraph(doc, "背景技术", bold=True, indent=False)
    for paragraph in spec["background"]:
        add_text_paragraph(doc, paragraph)

    add_text_paragraph(doc, "发明内容", bold=True, indent=False)
    invention = spec["invention_content"]
    for paragraph in invention["problem"]:
        add_text_paragraph(doc, paragraph)
    add_text_paragraph(doc, "本发明的技术解决方案是：")
    add_text_paragraph(doc, f"{data['title']}，包括以下步骤：")
    for step in method_steps(data["claims"][0]["text"]):
        add_text_paragraph(doc, step)
    for paragraph in invention["solution"]:
        add_text_paragraph(doc, paragraph.replace("动作执行后，系统将", "动作执行后，该方法将"))
    add_beneficial_effects(doc, invention["beneficial_effects"])

    add_text_paragraph(doc, "附图说明", bold=True, indent=False)
    for paragraph in spec["figure_descriptions"]:
        add_text_paragraph(doc, paragraph)

    add_text_paragraph(doc, "具体实施方式", bold=True, indent=False)
    add_text_paragraph(doc, "下面结合附图详细说明本发明的优选实施例。")
    add_text_paragraph(doc, f"{data['title']}，如图1所示，包括以下步骤：")
    for step in method_steps(data["claims"][0]["text"]):
        add_text_paragraph(doc, step)
    equation = spec["equations"][0]["expression"] if spec.get("equations") else None
    for embodiment in spec["embodiments"]:
        add_text_paragraph(doc, embodiment_lead(embodiment["heading"]))
        for paragraph in embodiment["paragraphs"]:
            if paragraph.strip().startswith("\\[") and equation:
                add_equation(doc, equation)
            else:
                add_text_paragraph(doc, paragraph)
    add_text_paragraph(
        doc,
        "以上为本发明的优选实施方式。在不脱离本发明方法技术原理的情况下，对步骤顺序、数据表示或动作评价实现作出的等同替换，应由权利要求所限定的保护范围确定。",
    )


def image_size(path, max_width, max_height):
    with Image.open(path) as image:
        width, height = image.size
    ratio = min(max_width / width, max_height / height)
    return Inches(width * ratio / 96), Inches(height * ratio / 96)


def add_figures(doc, figures_dir, descriptions):
    images = sorted(figures_dir.glob("figure-*.png"), key=lambda p: int(p.stem.split("-")[-1]))
    if len(images) != 5:
        raise ValueError(f"Expected 5 figures, found {len(images)} in {figures_dir}")
    for index, image_path in enumerate(images, start=1):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        max_height = 3.25 if index < 5 else 6.5
        width, height = image_size(image_path, 520, max_height * 96)
        inline_shape = paragraph.add_run().add_picture(str(image_path), width=width, height=height)
        description = descriptions[index - 1] if index <= len(descriptions) else f"图{index}"
        inline_shape._inline.docPr.set("title", f"图{index}")
        inline_shape._inline.docPr.set("descr", description)
        add_text_paragraph(doc, f"图{index}", center=True, indent=False)
        if index in {2, 4}:
            doc.add_page_break()


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = BODY_SIZE
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)


def build(args):
    data = json.loads(args.source_json.read_text(encoding="utf-8"))
    if [claim["number"] for claim in data["claims"]] != list(range(1, 10)):
        raise ValueError("The standard-format package requires method claims 1 through 9 only.")
    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0], "摘要")
    if doc.paragraphs:
        set_text_paragraph(doc.paragraphs[0], "摘要", center=True, indent=False)
    else:
        add_text_paragraph(doc, "摘要", center=True, indent=False)
    add_text_paragraph(doc, data["abstract"])

    add_new_section(doc, "权利要求书")
    add_text_paragraph(doc, "权利要求书", center=True, indent=False)
    add_claims(doc, data["claims"])

    add_new_section(doc, "说明书")
    add_text_paragraph(doc, "说明书", center=True, indent=False)
    add_text_paragraph(doc, data["title"], center=True, indent=False)
    add_specification(doc, data)

    add_new_section(doc, "说明书附图")
    add_text_paragraph(doc, "说明书附图", center=True, indent=False)
    add_figures(doc, args.figures_dir, data["specification"]["figure_descriptions"])

    doc.core_properties.title = data["title"]
    doc.core_properties.subject = "中国发明专利方法稿"
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.output)


def parse_args():
    parser = argparse.ArgumentParser(description="Build a ZJU-style method-only Chinese patent DOCX.")
    parser.add_argument("--source-json", type=Path, required=True)
    parser.add_argument("--figures-dir", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    return parser.parse_args()


if __name__ == "__main__":
    build(parse_args())
