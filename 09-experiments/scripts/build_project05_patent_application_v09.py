from __future__ import annotations

import argparse
import hashlib
import json
import re
import shutil
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Cm, Inches, Pt
from PIL import Image


ROOT = Path(__file__).resolve().parents[2]
WRITING = ROOT / "08-writing"
DEFAULT_SOURCE = WRITING / "patent-application-text-v0.9-zju-reference-20260715.md"
DEFAULT_REFERENCE = Path(
    "C:/Users/35393/Desktop/浙大项目/年度报告/"
    "一种融合根源语义信息的高级持续性威胁检测与溯源方法-9-24-reviewed.docx"
)
DEFAULT_PACKAGE = WRITING / "patent-package-v0.9-zju-reference"
DEFAULT_FIGURES = DEFAULT_PACKAGE / "Project05_调查取证动作规划方法-figures"
DEFAULT_OUTPUT = DEFAULT_PACKAGE / "Project05_调查取证动作规划方法-浙大参考格式.docx"
DEFAULT_AUDIT = DEFAULT_PACKAGE / "build-audit-v0.9.json"
EXPECTED_REFERENCE_SHA256 = "D5D1E41421F973BDBCEE432A45C5CE0F5B7635BFCE4DE609200BEB10477ADBB1"

BODY_SIZE = Pt(14)
HEADER_SIZE = Pt(9)
FIRST_LINE = Pt(28)
FIGURE_WIDTH = Inches(5.75)
FIGURE_MAX_HEIGHT = Inches(7.75)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


def h1_sections(text: str) -> dict[str, str]:
    matches = list(re.finditer(r"(?m)^# ([^#\r\n].+?)\s*$", text))
    output: dict[str, str] = {}
    for index, match in enumerate(matches):
        end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
        output[match.group(1).strip()] = text[match.end() : end].strip()
    return output


def h2_sections(text: str) -> tuple[list[str], dict[str, str]]:
    matches = list(re.finditer(r"(?m)^## ([^#\r\n].+?)\s*$", text))
    order: list[str] = []
    output: dict[str, str] = {}
    for index, match in enumerate(matches):
        end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
        name = match.group(1).strip()
        order.append(name)
        output[name] = text[match.end() : end].strip()
    return order, output


def blocks(text: str) -> list[str]:
    return [
        re.sub(r"\s*\n\s*", " ", block).strip()
        for block in re.split(r"\n\s*\n", text)
        if block.strip()
    ]


def parse_source(path: Path) -> dict:
    text = path.read_text(encoding="utf-8-sig")
    sections = h1_sections(text)
    required = {"摘要", "权利要求书", "说明书", "说明书附图"}
    if set(sections) != required:
        raise ValueError(f"Expected exactly {sorted(required)} H1 sections, found {sorted(sections)}")

    abstract_blocks = blocks(sections["摘要"])
    if len(abstract_blocks) != 1:
        raise ValueError(f"摘要 must contain one paragraph, found {len(abstract_blocks)}")

    claims: list[dict[str, object]] = []
    for block in blocks(sections["权利要求书"]):
        match = re.fullmatch(r"(\d+)\.\s+(.+)", block)
        if match is None:
            raise ValueError(f"Invalid claim paragraph: {block[:80]}")
        claims.append({"number": int(match.group(1)), "text": match.group(2).strip()})
    if [item["number"] for item in claims] != list(range(1, 12)):
        raise ValueError("The v0.9 application must contain method claims 1 through 11.")

    order, specification = h2_sections(sections["说明书"])
    expected_order = ["一种面向不完整安全证据的调查取证动作规划方法", "技术领域", "背景技术", "发明内容", "附图说明", "具体实施方式"]
    if order != expected_order:
        raise ValueError(f"Unexpected specification section order: {order}")
    title = order[0]
    content = {name: blocks(specification[name]) for name in order[1:]}
    figure_descriptions = content["附图说明"]
    if len(figure_descriptions) != 5:
        raise ValueError(f"Expected 5 figure descriptions, found {len(figure_descriptions)}")
    return {
        "raw": text,
        "title": title,
        "abstract": abstract_blocks[0],
        "claims": claims,
        "specification": content,
        "figure_descriptions": figure_descriptions,
    }


def set_run_font(run, *, size=BODY_SIZE, bold: bool = False):
    run.font.name = "Times New Roman"
    run.font.size = size
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "宋体")


def add_paragraph(
    doc: Document,
    value: str,
    *,
    center: bool = False,
    bold: bool = False,
    indent: bool = True,
    keep_with_next: bool | None = None,
):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(value)
    set_run_font(run, bold=bold)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = paragraph.paragraph_format
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.first_line_indent = FIRST_LINE if indent else Pt(0)
    fmt.keep_with_next = keep_with_next
    return paragraph


def remove_references(sect_pr):
    for child in list(sect_pr):
        if child.tag in {qn("w:headerReference"), qn("w:footerReference")}:
            sect_pr.remove(child)


def reset_reference_body(doc: Document):
    body = doc._element.body
    final_sect_pr = body.sectPr
    if final_sect_pr is None:
        raise ValueError("Reference DOCX has no final section properties.")
    for child in list(body):
        if child is not final_sect_pr:
            body.remove(child)
    remove_references(final_sect_pr)
    for relation_id, relation in list(doc.part.rels.items()):
        if relation.reltype in {RT.HEADER, RT.FOOTER, RT.IMAGE}:
            doc.part.drop_rel(relation_id)


def configure_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = BODY_SIZE
    normal.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    header = doc.styles["Header"]
    header.font.name = "Times New Roman"
    header.font.size = HEADER_SIZE
    header.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def configure_section(section, header_text: str):
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3.175)
    section.right_margin = Cm(3.175)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.header_distance = Cm(1.50)
    section.footer_distance = Cm(1.75)
    section.different_first_page_header_footer = False
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False

    header_paragraph = section.header.paragraphs[0]
    clear_paragraph(header_paragraph)
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_paragraph.paragraph_format.space_before = Pt(0)
    header_paragraph.paragraph_format.space_after = Pt(0)
    header_paragraph.paragraph_format.first_line_indent = Pt(0)
    run = header_paragraph.add_run(header_text)
    set_run_font(run, size=HEADER_SIZE)
    ppr = header_paragraph._p.get_or_add_pPr()
    border = ppr.find(qn("w:pBdr"))
    if border is not None:
        ppr.remove(border)

    for paragraph in list(section.footer.paragraphs):
        clear_paragraph(paragraph)


def add_section(doc: Document, header_text: str):
    section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    configure_section(section, header_text)
    return section


def add_image(doc: Document, path: Path, description: str):
    with Image.open(path) as image:
        width_px, height_px = image.size
    ratio = min(FIGURE_WIDTH.inches / width_px, FIGURE_MAX_HEIGHT.inches / height_px)
    width = Inches(width_px * ratio)
    height = Inches(height_px * ratio)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    shape = paragraph.add_run().add_picture(str(path), width=width, height=height)
    shape._inline.docPr.set("title", path.stem)
    shape._inline.docPr.set("descr", description)


def build_document(data: dict, reference: Path, figures_dir: Path, output: Path):
    output.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(reference, output)
    doc = Document(output)
    reset_reference_body(doc)
    configure_styles(doc)

    configure_section(doc.sections[0], "摘要")
    add_paragraph(doc, "摘要", center=True, indent=False)
    add_paragraph(doc, data["abstract"])

    add_section(doc, "权利要求书")
    add_paragraph(doc, "权利要求书", center=True, indent=False)
    for claim in data["claims"]:
        add_paragraph(doc, f"{claim['number']}. {claim['text']}")

    add_section(doc, "说明书")
    add_paragraph(doc, "说明书", center=True, indent=False)
    add_paragraph(doc, data["title"], center=True, indent=False)
    for section_name in ["技术领域", "背景技术", "发明内容", "附图说明", "具体实施方式"]:
        add_paragraph(doc, section_name, bold=True, indent=False, keep_with_next=True)
        for paragraph in data["specification"][section_name]:
            add_paragraph(doc, paragraph)

    add_section(doc, "说明书附图")
    add_paragraph(doc, "说明书附图", center=True, indent=False)
    images = sorted(figures_dir.glob("figure-*.png"), key=lambda path: int(path.stem.split("-")[-1]))
    if [path.name for path in images] != [f"figure-{index}.png" for index in range(1, 6)]:
        raise ValueError(f"Expected figure-1.png through figure-5.png in {figures_dir}")
    for index, image_path in enumerate(images, start=1):
        if index > 1:
            doc.add_page_break()
        add_paragraph(doc, f"图{index}", center=True, indent=False, keep_with_next=True)
        add_image(doc, image_path, data["figure_descriptions"][index - 1])

    doc.core_properties.title = data["title"]
    doc.core_properties.subject = "中国发明专利申请文本（浙大参考格式）"
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    doc.core_properties.comments = ""
    doc.save(output)


def cjk_count(value: str) -> int:
    return len(re.findall(r"[\u3400-\u4dbf\u4e00-\u9fff\uf900-\ufaff]", value))


def audit_document(data: dict, reference: Path, output: Path) -> dict:
    doc = Document(output)
    body_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    claims_text = "\n".join(
        paragraph.text for paragraph in doc.paragraphs if re.match(r"^\d+\.\s", paragraph.text)
    )
    independent = next(paragraph.text for paragraph in doc.paragraphs if paragraph.text.startswith("1. "))
    headers = [section.header.paragraphs[0].text for section in doc.sections]
    footers = ["".join(paragraph.text for paragraph in section.footer.paragraphs) for section in doc.sections]
    forbidden_body = [
        "all_experiments_complete",
        "P/E/F/C",
        "提交红线",
        "08-writing/",
        "09-experiments/",
        "```",
        "一种融合根源语义信息的高级持续性威胁检测与溯源方法",
        "一种基于系统溯源图的高级持续威胁检测与识别方法",
    ]
    forbidden_independent = ["LLM", "DQN", "XGBoost", "DARPA", "OpTC", "C07", "C12"]
    package_text = ""
    package_entries: list[str] = []
    fields = 0
    tracked_changes = 0
    with ZipFile(output) as archive:
        package_entries = sorted(archive.namelist())
        for name in package_entries:
            if name.endswith(".xml") or name.endswith(".rels"):
                payload = archive.read(name).decode("utf-8", errors="ignore")
                package_text += payload
                fields += payload.count("<w:fldChar") + payload.count("<w:instrText")
                tracked_changes += len(re.findall(r"<w:(?:ins|del)(?:\s|>)", payload))

    checks = {
        "reference_sha256": sha256(reference) == EXPECTED_REFERENCE_SHA256,
        "four_sections": len(doc.sections) == 4,
        "section_headers": headers == ["摘要", "权利要求书", "说明书", "说明书附图"],
        "empty_footers": footers == ["", "", "", ""],
        "a4_geometry": all(
            abs(section.page_width.cm - 21.0) < 0.02
            and abs(section.page_height.cm - 29.7) < 0.02
            and abs(section.left_margin.cm - 3.175) < 0.02
            and abs(section.right_margin.cm - 3.175) < 0.02
            and abs(section.top_margin.cm - 2.54) < 0.02
            and abs(section.bottom_margin.cm - 2.54) < 0.02
            for section in doc.sections
        ),
        "five_inline_figures": len(doc.inline_shapes) == 5,
        "no_tables": len(doc.tables) == 0,
        "eleven_method_claims": len(re.findall(r"(?m)^\d+\.\s", claims_text)) == 11
        and "一种电子设备" not in claims_text
        and "计算机可读存储介质" not in claims_text,
        "independent_claim_generic": not any(term in independent for term in forbidden_independent),
        "no_forbidden_body_content": not any(term in body_text for term in forbidden_body),
        "no_reference_content_in_package": not any(term in package_text for term in forbidden_body[-2:]),
        "no_fields": fields == 0,
        "no_tracked_changes": tracked_changes == 0,
        "no_comments_part": not any("comments" in name.lower() for name in package_entries),
    }
    if not all(checks.values()):
        failed = [name for name, passed in checks.items() if not passed]
        raise ValueError(f"Patent DOCX audit failed: {failed}")
    visible_nonspace = len(re.sub(r"\s+", "", body_text))
    return {
        "status": "PASS",
        "source": str(DEFAULT_SOURCE),
        "reference": str(reference),
        "reference_sha256": sha256(reference),
        "output": str(output),
        "output_sha256": sha256(output),
        "visible_nonspace_characters": visible_nonspace,
        "visible_cjk_characters": cjk_count(body_text),
        "claim_count": 11,
        "section_count": len(doc.sections),
        "figure_count": len(doc.inline_shapes),
        "checks": checks,
    }


def main():
    parser = argparse.ArgumentParser(description="Build the Project05 v0.9 ZJU-reference patent DOCX.")
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--reference", type=Path, default=DEFAULT_REFERENCE)
    parser.add_argument("--figures-dir", type=Path, default=DEFAULT_FIGURES)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    parser.add_argument("--audit", type=Path, default=DEFAULT_AUDIT)
    args = parser.parse_args()

    if not args.reference.exists():
        raise FileNotFoundError(args.reference)
    actual_reference_hash = sha256(args.reference)
    if actual_reference_hash != EXPECTED_REFERENCE_SHA256:
        raise ValueError(
            f"Reference SHA-256 mismatch: expected {EXPECTED_REFERENCE_SHA256}, got {actual_reference_hash}"
        )
    data = parse_source(args.source)
    build_document(data, args.reference, args.figures_dir, args.output)
    audit = audit_document(data, args.reference, args.output)
    args.audit.parent.mkdir(parents=True, exist_ok=True)
    args.audit.write_text(json.dumps(audit, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(audit, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
