import json
import re
import unittest
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[2]
WRITING = ROOT / "08-writing"
WORK = WRITING / "patent-work"
PACKAGE = WRITING / "patent-package-v0.7-zju-format"
MAIN = WRITING / "patent-main-draft-v0.7-zju-format-20260714.md"
STRUCTURED = WORK / "07-structured-draft-v0.7.json"
CLAIMS = WORK / "06-claims-v0.7.txt"
DOCX = PACKAGE / "Project05_归因取证动作规划方法-浙大标准格式.docx"
AUTHORITY = WRITING / "AUTHORITATIVE-DOCUMENTS-20260713.md"


class PatentV07ZjuFormatTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.data = json.loads(STRUCTURED.read_text(encoding="utf-8"))
        cls.main = MAIN.read_text(encoding="utf-8")
        cls.claims = CLAIMS.read_text(encoding="utf-8-sig")
        cls.doc = Document(DOCX) if DOCX.exists() else None

    @unittest.skipUnless(DOCX.exists(), "generated DOCX is not tracked in Git")
    def test_package_contains_the_final_docx(self):
        self.assertTrue(DOCX.is_file())
        self.assertGreater(DOCX.stat().st_size, 100_000)

    @unittest.skipUnless(DOCX.exists(), "generated DOCX is not tracked in Git")
    def test_four_reference_style_sections(self):
        self.assertEqual(len(self.doc.sections), 4)
        headers = [section.header.paragraphs[0].text for section in self.doc.sections]
        self.assertEqual(headers, ["摘要", "权利要求书", "说明书", "说明书附图"])
        for section in self.doc.sections:
            self.assertAlmostEqual(section.page_width.cm, 21.0, places=1)
            self.assertAlmostEqual(section.page_height.cm, 29.7, places=1)
            self.assertAlmostEqual(section.left_margin.cm, 3.175, places=2)
            self.assertAlmostEqual(section.right_margin.cm, 3.175, places=2)
            self.assertAlmostEqual(section.top_margin.cm, 2.54, places=2)
            self.assertAlmostEqual(section.bottom_margin.cm, 2.54, places=2)

    @unittest.skipUnless(DOCX.exists(), "generated DOCX is not tracked in Git")
    def test_body_section_order_and_figures(self):
        body = [paragraph.text for paragraph in self.doc.paragraphs if paragraph.text]
        positions = [body.index(title) for title in ["摘要", "权利要求书", "说明书", "说明书附图"]]
        self.assertEqual(positions, sorted(positions))
        self.assertEqual(len(self.doc.inline_shapes), 5)

    def test_claim_architecture_remains_method_only(self):
        self.assertEqual([claim["number"] for claim in self.data["claims"]], list(range(1, 10)))
        self.assertEqual(re.findall(r"(?m)^(\d+)\.\s", self.claims), [str(i) for i in range(1, 10)])
        forbidden = ["一种电子设备", "一种计算机可读存储介质", "系统权利要求"]
        for phrase in forbidden:
            self.assertNotIn(phrase, self.claims)
        self.assertIn("由计算设备执行", self.data["claims"][0]["text"])

    def test_no_markdown_code_markers_in_formal_content(self):
        self.assertNotIn("`", self.main)
        self.assertNotIn("`", json.dumps(self.data, ensure_ascii=False))

    @unittest.skipUnless(DOCX.exists(), "generated DOCX is not tracked in Git")
    def test_headers_have_bottom_rule(self):
        for section in self.doc.sections:
            ppr = section.header.paragraphs[0]._p.pPr
            bottom = ppr.find(qn("w:pBdr")).find(qn("w:bottom"))
            self.assertEqual(bottom.get(qn("w:val")), "single")

    def test_authority_points_to_v07(self):
        authority = AUTHORITY.read_text(encoding="utf-8")
        self.assertIn("patent-main-draft-v0.7-zju-format-20260714.md", authority)
        self.assertIn("patent-package-v0.7-zju-format/", authority)
        self.assertIn("系统、设备及存储介质权利要求保持删除状态", authority)

    def test_validation_report_passed(self):
        report = (WORK / "08-validation-report-v0.7.txt").read_text(encoding="utf-8")
        self.assertIn("PASS", report)


if __name__ == "__main__":
    unittest.main()
