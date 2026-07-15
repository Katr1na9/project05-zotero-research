from __future__ import annotations

import argparse
import json
import re
import zipfile
from collections import Counter
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


def value_or_none(value):
    if value is None:
        return None
    if hasattr(value, "pt"):
        return round(value.pt, 3)
    return str(value)


def paragraph_record(index, paragraph):
    pf = paragraph.paragraph_format
    runs = []
    for run in paragraph.runs:
        text = run.text
        if not text:
            continue
        rpr = run._element.rPr
        fonts = rpr.rFonts if rpr is not None else None
        east_asia = fonts.get(qn("w:eastAsia")) if fonts is not None else None
        runs.append(
            {
                "text": text,
                "font": run.font.name,
                "east_asia_font": east_asia,
                "size_pt": value_or_none(run.font.size),
                "bold": run.bold,
                "italic": run.italic,
                "underline": bool(run.underline) if run.underline is not None else None,
            }
        )
    ppr = paragraph._element.pPr
    sect_break = bool(ppr is not None and ppr.sectPr is not None)
    return {
        "index": index,
        "text": paragraph.text,
        "style": paragraph.style.name if paragraph.style is not None else None,
        "alignment": str(paragraph.alignment),
        "left_indent_pt": value_or_none(pf.left_indent),
        "right_indent_pt": value_or_none(pf.right_indent),
        "first_line_indent_pt": value_or_none(pf.first_line_indent),
        "space_before_pt": value_or_none(pf.space_before),
        "space_after_pt": value_or_none(pf.space_after),
        "line_spacing": value_or_none(pf.line_spacing),
        "keep_with_next": pf.keep_with_next,
        "page_break_before": pf.page_break_before,
        "section_break_after": sect_break,
        "runs": runs,
    }


def read_app_properties(docx_path: Path):
    with zipfile.ZipFile(docx_path) as archive:
        xml = archive.read("docProps/app.xml").decode("utf-8", errors="replace")
        document_xml = archive.read("word/document.xml").decode(
            "utf-8", errors="replace"
        )
        part_names = sorted(archive.namelist())
    props = {}
    for name in ("Pages", "Words", "Characters", "Paragraphs", "Lines"):
        match = re.search(rf"<{name}>(.*?)</{name}>", xml)
        props[name.lower()] = int(match.group(1)) if match else None
    props["last_rendered_page_breaks"] = document_xml.count("lastRenderedPageBreak")
    props["tracked_insertions"] = document_xml.count("<w:ins")
    props["tracked_deletions"] = document_xml.count("<w:del")
    props["comment_parts"] = [name for name in part_names if "comment" in name.lower()]
    props["headers"] = [name for name in part_names if re.match(r"word/header\d+\.xml", name)]
    props["footers"] = [name for name in part_names if re.match(r"word/footer\d+\.xml", name)]
    return props


def style_record(style):
    font = style.font
    rpr = style._element.rPr
    fonts = rpr.rFonts if rpr is not None else None
    pf = style.paragraph_format
    return {
        "name": style.name,
        "type": str(style.type),
        "base_style": style.base_style.name if style.base_style is not None else None,
        "font_name": font.name,
        "ascii_font": fonts.get(qn("w:ascii")) if fonts is not None else None,
        "hansi_font": fonts.get(qn("w:hAnsi")) if fonts is not None else None,
        "east_asia_font": fonts.get(qn("w:eastAsia")) if fonts is not None else None,
        "size_pt": value_or_none(font.size),
        "bold": font.bold,
        "italic": font.italic,
        "alignment": str(pf.alignment),
        "left_indent_pt": value_or_none(pf.left_indent),
        "first_line_indent_pt": value_or_none(pf.first_line_indent),
        "space_before_pt": value_or_none(pf.space_before),
        "space_after_pt": value_or_none(pf.space_after),
        "line_spacing": value_or_none(pf.line_spacing),
    }


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("docx", type=Path)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()

    document = Document(args.docx)
    paragraphs = [
        paragraph_record(index, paragraph)
        for index, paragraph in enumerate(document.paragraphs)
    ]
    nonempty = [item for item in paragraphs if item["text"].strip()]
    text = "\n".join(item["text"] for item in nonempty)

    section_records = []
    for index, section in enumerate(document.sections, start=1):
        section_records.append(
            {
                "index": index,
                "start_type": str(section.start_type),
                "page_width_in": round(section.page_width.inches, 3),
                "page_height_in": round(section.page_height.inches, 3),
                "left_margin_in": round(section.left_margin.inches, 3),
                "right_margin_in": round(section.right_margin.inches, 3),
                "top_margin_in": round(section.top_margin.inches, 3),
                "bottom_margin_in": round(section.bottom_margin.inches, 3),
                "header_distance_in": round(section.header_distance.inches, 3),
                "footer_distance_in": round(section.footer_distance.inches, 3),
                "header_text": " | ".join(
                    p.text for p in section.header.paragraphs if p.text.strip()
                ),
                "header_paragraphs": [
                    paragraph_record(i, p)
                    for i, p in enumerate(section.header.paragraphs)
                    if p.text.strip()
                ],
                "footer_text": " | ".join(
                    p.text for p in section.footer.paragraphs if p.text.strip()
                ),
                "footer_paragraphs": [
                    paragraph_record(i, p)
                    for i, p in enumerate(section.footer.paragraphs)
                    if p.text.strip()
                ],
            }
        )

    data = {
        "path": str(args.docx.resolve()),
        "app_properties": read_app_properties(args.docx),
        "counts": {
            "paragraphs": len(paragraphs),
            "nonempty_paragraphs": len(nonempty),
            "characters_no_whitespace": len(re.sub(r"\s+", "", text)),
            "cjk_characters": len(re.findall(r"[\u3400-\u9fff]", text)),
            "tables": len(document.tables),
            "inline_shapes": len(document.inline_shapes),
            "sections": len(document.sections),
        },
        "style_counts_nonempty": Counter(item["style"] for item in nonempty),
        "used_style_definitions": [
            style_record(document.styles[name])
            for name in sorted({item["style"] for item in nonempty})
            if name in document.styles
        ],
        "sections": section_records,
        "paragraphs": paragraphs,
    }
    args.output.write_text(
        json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8"
    )


if __name__ == "__main__":
    main()
